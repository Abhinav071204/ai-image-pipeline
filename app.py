import os
import io
import time
import threading
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import requests
from docx import Document
from docx.shared import Inches
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload

KEY_PATH = r"C:\Users\Riyak\Downloads\prompt-image-pipeline-9e02923b14b1.json"
LOCAL_OUTPUT = r"C:\ImagePipeline\output"
FOLDERS = {
    "inbox":  "1CyNIvtTvIGg9ohKKdkjt8EEoKw2q6nUL",
    "images": "1HznEAs-SehbmTijjpymdyFZurpGQPhlv",
    "done":   "1Xq5jZD_b76LRaXIfo-HYhgQeKsct4ZID",
}

# Colors
DARK = "#1e1e2e"
SIDEBAR = "#2a2a3d"
ACCENT = "#7c6af7"
LIGHT_BG = "#f5f5f5"
WHITE = "#ffffff"
TEXT_DARK = "#1e1e2e"
TEXT_LIGHT = "#ffffff"
TEXT_MUTED = "#888888"
SUCCESS = "#4caf50"
WARNING = "#ff9800"
ERROR = "#f44336"

def get_drive_service():
    creds = service_account.Credentials.from_service_account_file(
        KEY_PATH, scopes=["https://www.googleapis.com/auth/drive"]
    )
    return build("drive", "v3", credentials=creds)

def get_files_in_folder(service, folder_id):
    query = "'" + folder_id + "' in parents and trashed=false"
    results = service.files().list(q=query, fields="files(id, name, mimeType)").execute()
    return results.get("files", [])

def extract_prompts(docx_buffer):
    doc = Document(docx_buffer)
    prompts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            prompts.append(text)
    return prompts

def generate_image(prompt):
    safe_prompt = requests.utils.quote(prompt)
    url = "https://image.pollinations.ai/prompt/" + safe_prompt + "?width=512&height=512&nologo=true"
    response = requests.get(url, timeout=60)
    if response.status_code == 200:
        return response.content
    else:
        raise Exception("Image generation failed")

def download_file(service, file_id):
    request = service.files().get_media(fileId=file_id)
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buffer.seek(0)
    return buffer

def embed_and_upload(service, docx_buffer, prompts_and_images, file_id):
    doc = Document(docx_buffer)
    for i, (prompt, image_bytes) in enumerate(prompts_and_images):
        doc.add_paragraph("")
        doc.add_paragraph("Prompt " + str(i+1) + ": " + prompt)
        doc.add_paragraph("--- Generated Image ---")
        doc.add_picture(io.BytesIO(image_bytes), width=Inches(4.0))
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    media = MediaIoBaseUpload(output_buffer, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    service.files().update(fileId=file_id, media_body=media).execute()

def move_to_done(service, file_id):
    file = service.files().get(fileId=file_id, fields="parents").execute()
    previous_parents = ",".join(file.get("parents", []))
    service.files().update(fileId=file_id, addParents=FOLDERS["done"], removeParents=previous_parents, fields="id, parents").execute()

class PipelineApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AI Image Pipeline")
        self.root.geometry("900x650")
        self.root.configure(bg=DARK)
        self.root.resizable(True, True)
        self.images_refs = []
        self.build_ui()

    def build_ui(self):
        # Sidebar
        sidebar = tk.Frame(self.root, bg=SIDEBAR, width=200)
        sidebar.pack(side="left", fill="y")
        sidebar.pack_propagate(False)

        tk.Label(sidebar, text="AI Pipeline", bg=SIDEBAR, fg=TEXT_LIGHT,
                 font=("Arial", 16, "bold")).pack(pady=30)

        buttons = [
            ("Upload & Run", self.show_upload),
            ("Image Gallery", self.show_gallery),
            ("File Status", self.show_status),
        ]
        for label, cmd in buttons:
            btn = tk.Button(sidebar, text=label, bg=ACCENT, fg=TEXT_LIGHT,
                            font=("Arial", 11), bd=0, padx=20, pady=10,
                            cursor="hand2", activebackground="#6a58e0",
                            activeforeground=TEXT_LIGHT, command=cmd)
            btn.pack(fill="x", padx=20, pady=5)

        tk.Label(sidebar, text="v1.0", bg=SIDEBAR, fg=TEXT_MUTED,
                 font=("Arial", 9)).pack(side="bottom", pady=10)

        # Main content area
        self.content = tk.Frame(self.root, bg=LIGHT_BG)
        self.content.pack(side="right", fill="both", expand=True)

        self.show_upload()

    def clear_content(self):
        for widget in self.content.winfo_children():
            widget.destroy()

    def show_upload(self):
        self.clear_content()

        tk.Label(self.content, text="Upload & Run Pipeline", bg=LIGHT_BG,
                 fg=TEXT_DARK, font=("Arial", 18, "bold")).pack(pady=30)

        tk.Label(self.content, text="Select a Word document with prompts (one per line)",
                 bg=LIGHT_BG, fg=TEXT_MUTED, font=("Arial", 11)).pack()

        self.file_label = tk.Label(self.content, text="No file selected",
                                    bg=WHITE, fg=TEXT_MUTED, font=("Arial", 10),
                                    relief="groove", padx=20, pady=10, width=40)
        self.file_label.pack(pady=20)

        tk.Button(self.content, text="Browse File", bg=ACCENT, fg=TEXT_LIGHT,
                  font=("Arial", 11), bd=0, padx=20, pady=8, cursor="hand2",
                  command=self.browse_file).pack()

        self.run_btn = tk.Button(self.content, text="Run Pipeline", bg=SUCCESS,
                                  fg=TEXT_LIGHT, font=("Arial", 12, "bold"),
                                  bd=0, padx=30, pady=10, cursor="hand2",
                                  command=self.run_pipeline, state="disabled")
        self.run_btn.pack(pady=15)

        self.log_box = tk.Text(self.content, height=12, bg=DARK, fg="#00ff99",
                                font=("Courier", 10), bd=0, padx=10, pady=10,
                                state="disabled")
        self.log_box.pack(fill="both", expand=True, padx=20, pady=10)

    def browse_file(self):
        path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if path:
            self.selected_file = path
            self.file_label.config(text=os.path.basename(path), fg=TEXT_DARK)
            self.run_btn.config(state="normal")

    def log(self, message):
        self.log_box.config(state="normal")
        self.log_box.insert("end", message + "\n")
        self.log_box.see("end")
        self.log_box.config(state="disabled")

    def run_pipeline(self):
        self.run_btn.config(state="disabled")
        self.log_box.config(state="normal")
        self.log_box.delete("1.0", "end")
        self.log_box.config(state="disabled")
        thread = threading.Thread(target=self.pipeline_thread)
        thread.daemon = True
        thread.start()

    def pipeline_thread(self):
        try:
            self.log("Connecting to Google Drive...")
            service = get_drive_service()
            self.log("Connected!")

            # Upload file to inbox
            from googleapiclient.http import MediaFileUpload
            file_metadata = {"name": os.path.basename(self.selected_file), "parents": [FOLDERS["inbox"]]}
            media = MediaFileUpload(self.selected_file, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            uploaded = service.files().create(body=file_metadata, media_body=media, fields="id, name").execute()
            self.log("Uploaded: " + uploaded["name"])

            file_id = uploaded["id"]

            # Download and read
            docx_buffer = download_file(service, file_id)
            prompts = extract_prompts(docx_buffer)
            self.log("Found " + str(len(prompts)) + " prompt(s)")

            os.makedirs(LOCAL_OUTPUT, exist_ok=True)
            prompts_and_images = []

            for i, prompt in enumerate(prompts):
                self.log("Generating image " + str(i+1) + ": " + prompt)
                image_bytes = generate_image(prompt)
                filename = "prompt" + str(i+1) + "_" + str(int(time.time())) + ".png"
                local_path = os.path.join(LOCAL_OUTPUT, filename)
                with open(local_path, "wb") as f:
                    f.write(image_bytes)
                prompts_and_images.append((prompt, image_bytes))
                self.log("Image saved: " + filename)
                time.sleep(2)

            docx_buffer.seek(0)
            embed_and_upload(service, docx_buffer, prompts_and_images, file_id)
            self.log("Images embedded into document!")

            move_to_done(service, file_id)
            self.log("Document moved to done folder!")
            self.log("Pipeline complete!")

        except Exception as e:
            self.log("Error: " + str(e))
        finally:
            self.run_btn.config(state="normal")

    def show_gallery(self):
        self.clear_content()
        tk.Label(self.content, text="Image Gallery", bg=LIGHT_BG,
                 fg=TEXT_DARK, font=("Arial", 18, "bold")).pack(pady=20)

        canvas = tk.Canvas(self.content, bg=LIGHT_BG, bd=0, highlightthickness=0)
        scrollbar = tk.Scrollbar(self.content, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        canvas.pack(fill="both", expand=True)

        frame = tk.Frame(canvas, bg=LIGHT_BG)
        canvas.create_window((0, 0), window=frame, anchor="nw")

        self.images_refs = []
        if os.path.exists(LOCAL_OUTPUT):
            images = [f for f in os.listdir(LOCAL_OUTPUT) if f.endswith(".png")]
            if not images:
                tk.Label(frame, text="No images yet. Run the pipeline first!",
                         bg=LIGHT_BG, fg=TEXT_MUTED, font=("Arial", 12)).pack(pady=40)
            else:
                col, row = 0, 0
                for img_file in images:
                    img_path = os.path.join(LOCAL_OUTPUT, img_file)
                    try:
                        img = Image.open(img_path)
                        img.thumbnail((180, 180))
                        photo = ImageTk.PhotoImage(img)
                        self.images_refs.append(photo)
                        cell = tk.Frame(frame, bg=WHITE, relief="groove", bd=1)
                        cell.grid(row=row, column=col, padx=10, pady=10)
                        tk.Label(cell, image=photo, bg=WHITE).pack(padx=5, pady=5)
                        tk.Label(cell, text=img_file[:20], bg=WHITE, fg=TEXT_MUTED,
                                 font=("Arial", 8)).pack()
                        col += 1
                        if col > 2:
                            col = 0
                            row += 1
                    except Exception:
                        pass
        else:
            tk.Label(frame, text="No images yet. Run the pipeline first!",
                     bg=LIGHT_BG, fg=TEXT_MUTED, font=("Arial", 12)).pack(pady=40)

        frame.update_idletasks()
        canvas.configure(scrollregion=canvas.bbox("all"))

    def show_status(self):
        self.clear_content()
        tk.Label(self.content, text="File Status", bg=LIGHT_BG,
                 fg=TEXT_DARK, font=("Arial", 18, "bold")).pack(pady=20)

        tk.Button(self.content, text="Refresh", bg=ACCENT, fg=TEXT_LIGHT,
                  font=("Arial", 10), bd=0, padx=15, pady=6, cursor="hand2",
                  command=self.show_status).pack()

        try:
            service = get_drive_service()
            for folder_name, folder_id in FOLDERS.items():
                files = get_files_in_folder(service, folder_id)
                color = {"inbox": WARNING, "images": ACCENT, "done": SUCCESS}[folder_name]
                section = tk.Frame(self.content, bg=LIGHT_BG)
                section.pack(fill="x", padx=30, pady=10)
                tk.Label(section, text=folder_name.upper() + " (" + str(len(files)) + " files)",
                         bg=color, fg=TEXT_LIGHT, font=("Arial", 11, "bold"),
                         padx=10, pady=5).pack(fill="x")
                if files:
                    for f in files:
                        tk.Label(section, text="  - " + f["name"], bg=WHITE,
                                 fg=TEXT_DARK, font=("Arial", 10),
                                 anchor="w", padx=10, pady=3).pack(fill="x")
                else:
                    tk.Label(section, text="  (empty)", bg=WHITE,
                             fg=TEXT_MUTED, font=("Arial", 10),
                             anchor="w", padx=10, pady=3).pack(fill="x")
        except Exception as e:
            tk.Label(self.content, text="Error: " + str(e),
                     bg=LIGHT_BG, fg=ERROR, font=("Arial", 10)).pack(pady=20)

if __name__ == "__main__":
    root = tk.Tk()
    app = PipelineApp(root)
    root.mainloop()