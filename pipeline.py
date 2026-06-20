import os
import io
import time
import requests
from docx import Document
from docx.shared import Inches
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from config import KEY_PATH, LOCAL_OUTPUT, FOLDERS

def get_drive_service():
    creds = service_account.Credentials.from_service_account_file(KEY_PATH, scopes=["https://www.googleapis.com/auth/drive"])
    return build("drive", "v3", credentials=creds)

def get_inbox_files(service):
    query = "'" + FOLDERS["inbox"] + "' in parents and trashed=false and mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'"
    results = service.files().list(q=query, fields="files(id, name)").execute()
    return results.get("files", [])

def download_file(service, file_id):
    request = service.files().get_media(fileId=file_id)
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buffer.seek(0)
    return buffer

def extract_prompts(docx_buffer):
    doc = Document(docx_buffer)
    prompts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            prompts.append(text)
    return prompts

def generate_image(prompt):
    print("Generating image for: " + prompt)
    safe_prompt = requests.utils.quote(prompt)
    url = "https://image.pollinations.ai/prompt/" + safe_prompt + "?width=512&height=512&nologo=true"
    response = requests.get(url, timeout=60)
    if response.status_code == 200:
        print("Image generated!")
        return response.content
    else:
        raise Exception("Image generation failed: " + str(response.status_code))

def save_image_locally(image_bytes, filename):
    os.makedirs(LOCAL_OUTPUT, exist_ok=True)
    local_path = os.path.join(LOCAL_OUTPUT, filename)
    with open(local_path, "wb") as f:
        f.write(image_bytes)
    print("Image saved: " + local_path)
    return local_path

def embed_images_and_upload(service, docx_buffer, prompts_and_images, file_id):
    print("Embedding all images into document...")
    doc = Document(docx_buffer)
    for i, (prompt, image_bytes) in enumerate(prompts_and_images):
        doc.add_paragraph("")
        doc.add_paragraph("Prompt " + str(i+1) + ": " + prompt)
        doc.add_paragraph("--- Generated Image ---")
        doc.add_picture(io.BytesIO(image_bytes), width=Inches(4.0))
        doc.add_paragraph("")
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    media = MediaIoBaseUpload(output_buffer, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    service.files().update(fileId=file_id, media_body=media).execute()
    print("Document updated with all images!")

def move_to_done(service, file_id):
    print("Moving document to done folder...")
    file = service.files().get(fileId=file_id, fields="parents").execute()
    previous_parents = ",".join(file.get("parents", []))
    service.files().update(fileId=file_id, addParents=FOLDERS["done"], removeParents=previous_parents, fields="id, parents").execute()
    print("Document moved to done!")

def run_pipeline():
    print("Starting pipeline...")
    service = get_drive_service()
    files = get_inbox_files(service)
    if not files:
        print("No .docx files found in inbox. Upload a file and try again.")
        return
    print("Found " + str(len(files)) + " file(s) in inbox.")
    for file in files:
        print("Processing: " + file["name"])
        docx_buffer = download_file(service, file["id"])
        prompts = extract_prompts(docx_buffer)
        if not prompts:
            print("No prompts found. Skipping.")
            continue
        print("Found " + str(len(prompts)) + " prompt(s) in document.")
        prompts_and_images = []
        for i, prompt in enumerate(prompts):
            image_bytes = generate_image(prompt)
            image_filename = file["name"].replace(".docx", "") + "_prompt" + str(i+1) + ".png"
            save_image_locally(image_bytes, image_filename)
            prompts_and_images.append((prompt, image_bytes))
            time.sleep(2)
        docx_buffer.seek(0)
        embed_images_and_upload(service, docx_buffer, prompts_and_images, file["id"])
        move_to_done(service, file["id"])
        print("Done processing: " + file["name"])
    print("Pipeline complete!")

if __name__ == "__main__":
    print("Script started!")
    try:
        run_pipeline()
    except Exception as e:
        print("Error: " + str(e))
        import traceback
        traceback.print_exc()